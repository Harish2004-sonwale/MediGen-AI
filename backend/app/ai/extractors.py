from dataclasses import dataclass, field
import os
from pathlib import Path
from typing import Any
import docx
import pypdf


@dataclass
class ExtractedDocument:
    """Structured representation of extracted document contents."""

    text: str
    page_count: int
    pages: list[tuple[int, str]] = field(default_factory=list)  # (page_number, page_text)
    metadata: dict[str, Any] = field(default_factory=dict)


def extract_pdf(file_path: str) -> ExtractedDocument:
    """Extract text page-by-page from PDF file using pypdf."""
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"PDF document file not found on disk: {file_path}")

    try:
        reader = pypdf.PdfReader(file_path)
    except Exception as exc:
        raise ValueError(f"Failed to read PDF document structure: {exc}") from exc

    total_pages = len(reader.pages)
    if total_pages == 0:
        raise ValueError("PDF document has 0 pages.")

    pages: list[tuple[int, str]] = []
    full_text_parts: list[str] = []

    for idx, page in enumerate(reader.pages, start=1):
        try:
            page_text = page.extract_text() or ""
        except Exception:
            page_text = ""
        pages.append((idx, page_text))
        if page_text.strip():
            full_text_parts.append(page_text.strip())

    full_text = "\n\n".join(full_text_parts).strip()

    if not full_text:
        from app.core.config import settings
        if settings.OCR_ENABLED:
            from app.ai.ocr import get_ocr_provider
            ocr_provider = get_ocr_provider()
            return ocr_provider.extract_text(file_path, ".pdf")

        raise ValueError(
            "Document contains no extractable text (OCR image-based PDFs planned for future OCR milestone)"
        )

    return ExtractedDocument(
        text=full_text,
        page_count=total_pages,
        pages=pages,
        metadata={"format": "pdf", "total_pages": total_pages},
    )


def extract_docx(file_path: str) -> ExtractedDocument:
    """Extract text from Word DOCX file using python-docx."""
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"DOCX document file not found on disk: {file_path}")

    try:
        doc = docx.Document(file_path)
    except Exception as exc:
        raise ValueError(f"Failed to read DOCX document structure: {exc}") from exc

    paragraphs_text = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    # Extract tables content if present
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs_text.append(row_text)

    full_text = "\n\n".join(paragraphs_text).strip()

    if not full_text:
        raise ValueError("Document contains no extractable text.")

    return ExtractedDocument(
        text=full_text,
        page_count=1,
        pages=[(1, full_text)],
        metadata={"format": "docx", "paragraph_count": len(paragraphs_text)},
    )


def extract_txt(file_path: str) -> ExtractedDocument:
    """Extract text from plaintext TXT file with UTF-8 and fallback encoding."""
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"TXT document file not found on disk: {file_path}")

    try:
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            content = f.read()
    except Exception as exc:
        raise ValueError(f"Failed to read TXT document: {exc}") from exc

    full_text = content.strip()

    if not full_text:
        raise ValueError("Document contains no extractable text.")

    return ExtractedDocument(
        text=full_text,
        page_count=1,
        pages=[(1, full_text)],
        metadata={"format": "txt", "char_count": len(full_text)},
    )


def extract_document_text(file_path: str, file_extension: str) -> ExtractedDocument:
    """Dispatch document text extraction based on file extension."""
    ext = file_extension.lower().strip()
    if not ext.startswith("."):
        ext = f".{ext}"

    if ext == ".pdf":
        return extract_pdf(file_path)
    elif ext == ".docx":
        return extract_docx(file_path)
    elif ext == ".txt":
        return extract_txt(file_path)
    else:
        raise ValueError(f"Unsupported file format '{ext}' for text extraction.")

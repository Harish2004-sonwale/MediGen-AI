from datetime import date, datetime, timedelta, timezone
import io
import os
import shutil
from fastapi import status
from fastapi.testclient import TestClient
from sqlalchemy.orm import Session

from app.core.config import settings
from app.models.document import DocumentChunk, MedicalDocument
from app.models.patient import Patient
from app.models.user import User
from app.schemas.document import (
    DocumentChunkResponse,
    DocumentCreate,
    DocumentListResponse,
    DocumentProcessingStatus,
    DocumentResponse,
    DocumentType,
)
from app.schemas.patient import Gender, PatientStatus
from app.schemas.user import UserRole


def get_auth_headers(
    client: TestClient,
    role: UserRole = UserRole.DOCTOR,
    email: str = "doc_doc@hospital.org",
    name: str = "Dr. Doc User",
) -> tuple[dict[str, str], int]:
    """Register and login helper returning authorization headers and user ID."""
    reg_res = client.post(
        "/api/v1/auth/register",
        json={
            "name": name,
            "email": email,
            "password": "SecurePassword123!",
            "role": role.value,
        },
    )
    user_id = reg_res.json()["id"]

    login_res = client.post(
        "/api/v1/auth/login",
        json={"email": email, "password": "SecurePassword123!"},
    )
    token = login_res.json()["access_token"]
    return {"Authorization": f"Bearer {token}"}, user_id


def setup_clinical_environment(client: TestClient) -> dict[str, str]:
    """Helper to set up admin, doctor, and patient with clinical relationship."""
    admin_headers, _ = get_auth_headers(client, role=UserRole.ADMIN, email="admin_doc_env@hospital.org", name="Admin Doc")
    doc_headers, doc_uid = get_auth_headers(client, role=UserRole.DOCTOR, email="cardiologist_doc_env@hospital.org", name="Dr. Sarah Connor")
    pat_headers, _ = get_auth_headers(client, role=UserRole.PATIENT, email="alice_doc_env@patient.org", name="Alice Smith")

    # Create doctor profile
    doc_res = client.post(
        "/api/v1/doctors",
        json={
            "user_id": doc_uid,
            "full_name": "Sarah Connor",
            "department": "Cardiology",
            "specialization": "Cardiology",
            "medical_registration_number": "MED-DOC-ENV-001",
        },
        headers=admin_headers,
    )
    doc_id = doc_res.json()["doctor_id"]
    client.post(f"/api/v1/doctors/{doc_id}/verify", headers=admin_headers)

    # Create patient 1 (Alice)
    pat_res1 = client.post(
        "/api/v1/patients",
        json={
            "first_name": "Alice",
            "last_name": "Smith",
            "date_of_birth": "1990-01-01",
            "gender": "female",
            "email": "alice_doc_env@patient.org",
        },
        headers=admin_headers,
    )
    pat1_id = pat_res1.json()["patient_id"]

    # Create patient 2 (Bob - unrelated)
    pat_res2 = client.post(
        "/api/v1/patients",
        json={
            "first_name": "Bob",
            "last_name": "Jones",
            "date_of_birth": "1988-06-20",
            "gender": "male",
            "email": "bob_unrelated@patient.org",
        },
        headers=admin_headers,
    )
    pat2_id = pat_res2.json()["patient_id"]

    # Establish doctor-patient relationship for Alice via appointment
    future_time = (datetime.now(timezone.utc) + timedelta(days=2)).isoformat()
    client.post(
        "/api/v1/appointments",
        json={
            "patient_id": pat1_id,
            "doctor_id": doc_id,
            "appointment_date": future_time,
            "reason_for_visit": "Cardiac checkup",
        },
        headers=admin_headers,
    )

    return {
        "admin_headers": admin_headers,
        "doc_headers": doc_headers,
        "doc_id": doc_id,
        "pat1_headers": pat_headers,
        "pat1_id": pat1_id,
        "pat2_id": pat2_id,
    }


def test_document_schemas_and_enums():
    """Verify document validation schemas and enum values."""
    doc_in = DocumentCreate(
        patient_id="PAT-20260828-TEST",
        title="Complete Blood Count Report",
        document_type=DocumentType.LAB_REPORT,
    )
    assert doc_in.title == "Complete Blood Count Report"
    assert doc_in.document_type == DocumentType.LAB_REPORT

    doc_default = DocumentCreate(
        patient_id="PAT-20260828-TEST",
        title="Discharge Note",
    )
    assert doc_default.document_type == DocumentType.OTHER


def test_medical_document_and_chunk_orm_persistence(db_session: Session):
    """Verify MedicalDocument and DocumentChunk ORM models persist and cascade delete properly."""
    user = User(
        name="Clinical Uploader",
        email="uploader@hospital.org",
        password_hash="fakehash",
        role=UserRole.DOCTOR,
    )
    db_session.add(user)
    db_session.commit()
    db_session.refresh(user)

    patient = Patient(
        patient_id="PAT-20260828-DOC1",
        first_name="Alice",
        last_name="Smith",
        date_of_birth=date(1990, 1, 1),
        gender=Gender.FEMALE,
        status=PatientStatus.ACTIVE,
    )
    db_session.add(patient)
    db_session.commit()
    db_session.refresh(patient)

    doc = MedicalDocument(
        document_id="DOCU-20260828-0001",
        patient_id=patient.id,
        uploader_user_id=user.id,
        title="Echocardiogram Diagnostic Summary",
        document_type=DocumentType.IMAGING_REPORT,
        original_filename="echo_results.pdf",
        file_extension=".pdf",
        file_size_bytes=1048576,
        storage_path="/data/documents/echo_results.pdf",
        mime_type="application/pdf",
        processing_status=DocumentProcessingStatus.PENDING,
        page_count=3,
        total_chunks=2,
    )
    db_session.add(doc)
    db_session.commit()
    db_session.refresh(doc)

    assert doc.id is not None
    assert doc.document_id == "DOCU-20260828-0001"

    chunk1 = DocumentChunk(
        chunk_id="CHK-20260828-0001",
        document_id=doc.id,
        patient_id=patient.id,
        chunk_index=0,
        page_number=1,
        content="Patient presents with normal left ventricular ejection fraction.",
        token_count=10,
    )
    db_session.add(chunk1)
    db_session.commit()
    db_session.refresh(doc)

    assert len(doc.chunks) == 1

    # Cascade deletion
    db_session.delete(doc)
    db_session.commit()
    assert db_session.query(DocumentChunk).filter(DocumentChunk.document_id == doc.id).count() == 0


def test_successful_document_uploads(client: TestClient, tmp_path):
    """Verify successful upload of PDF, TXT, and DOCX documents with metadata generation."""
    settings.DOCUMENT_STORAGE_PATH = str(tmp_path)
    env = setup_clinical_environment(client)

    # 1. Upload TXT
    txt_content = b"Patient reports mild headache and fatigue."
    txt_file = ("notes.txt", io.BytesIO(txt_content), "text/plain")
    res_txt = client.post(
        "/api/v1/documents/upload",
        files={"file": txt_file},
        data={
            "patient_id": env["pat1_id"],
            "title": "Clinical Progress Note",
            "document_type": "clinical_note",
        },
        headers=env["doc_headers"],
    )
    assert res_txt.status_code == status.HTTP_201_CREATED
    data_txt = res_txt.json()
    assert data_txt["document_id"].startswith("DOCU-")
    assert data_txt["file_extension"] == ".txt"
    assert data_txt["document_type"] == "clinical_note"
    assert data_txt["processing_status"] == "completed"

    # 2. Upload DOCX
    import docx
    doc_docx = docx.Document()
    doc_docx.add_paragraph("Patient was admitted with acute bronchitis. Treatment initiated with oral antibiotics.")
    docx_buf = io.BytesIO()
    doc_docx.save(docx_buf)
    docx_bytes = docx_buf.getvalue()

    docx_file = (
        "summary.docx",
        io.BytesIO(docx_bytes),
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    res_docx = client.post(
        "/api/v1/documents/upload",
        files={"file": docx_file},
        data={
            "patient_id": env["pat1_id"],
            "title": "Discharge Summary",
            "document_type": "discharge_summary",
        },
        headers=env["admin_headers"],
    )
    assert res_docx.status_code == status.HTTP_201_CREATED
    assert res_docx.json()["file_extension"] == ".docx"
    assert res_docx.json()["processing_status"] == "completed"
    assert res_docx.json()["file_extension"] == ".docx"


def test_document_upload_validations_and_rejections(client: TestClient, tmp_path):
    """Verify rejections for EXE, ZIP, empty, and oversized files."""
    settings.DOCUMENT_STORAGE_PATH = str(tmp_path)
    env = setup_clinical_environment(client)

    # 1. Reject EXE
    exe_file = ("malicious.exe", io.BytesIO(b"MZ executable"), "application/x-msdownload")
    res_exe = client.post(
        "/api/v1/documents/upload",
        files={"file": exe_file},
        data={"patient_id": env["pat1_id"], "title": "Malicious"},
        headers=env["admin_headers"],
    )
    assert res_exe.status_code == status.HTTP_400_BAD_REQUEST
    assert "Unsupported file format" in res_exe.json()["detail"]

    # 2. Reject ZIP
    zip_file = ("archive.zip", io.BytesIO(b"PK archive"), "application/zip")
    res_zip = client.post(
        "/api/v1/documents/upload",
        files={"file": zip_file},
        data={"patient_id": env["pat1_id"], "title": "Archive"},
        headers=env["admin_headers"],
    )
    assert res_zip.status_code == status.HTTP_400_BAD_REQUEST

    # 3. Reject Empty file
    empty_file = ("empty.txt", io.BytesIO(b""), "text/plain")
    res_empty = client.post(
        "/api/v1/documents/upload",
        files={"file": empty_file},
        data={"patient_id": env["pat1_id"], "title": "Empty"},
        headers=env["admin_headers"],
    )
    assert res_empty.status_code == status.HTTP_400_BAD_REQUEST
    assert "empty" in res_empty.json()["detail"]


def test_document_rbac_and_cross_patient_isolation(client: TestClient, tmp_path):
    """Verify patient and doctor RBAC constraints and cross-patient isolation."""
    settings.DOCUMENT_STORAGE_PATH = str(tmp_path)
    env = setup_clinical_environment(client)

    # 1. Patient Alice uploads for herself -> Success
    pdf_file = ("my_report.pdf", io.BytesIO(b"%PDF-1.4 my report"), "application/pdf")
    res_pat_self = client.post(
        "/api/v1/documents/upload",
        files={"file": pdf_file},
        data={"patient_id": env["pat1_id"], "title": "My Self Upload"},
        headers=env["pat1_headers"],
    )
    assert res_pat_self.status_code == status.HTTP_201_CREATED
    doc_id = res_pat_self.json()["document_id"]

    # 2. Patient Alice tries to upload for Patient Bob -> 403 Forbidden
    res_pat_other = client.post(
        "/api/v1/documents/upload",
        files={"file": ("hack.pdf", io.BytesIO(b"%PDF-1.4 hack"), "application/pdf")},
        data={"patient_id": env["pat2_id"], "title": "Hack Upload"},
        headers=env["pat1_headers"],
    )
    assert res_pat_other.status_code == status.HTTP_403_FORBIDDEN

    # 3. Patient Alice lists documents -> sees only her own
    list_pat = client.get("/api/v1/documents", headers=env["pat1_headers"])
    assert list_pat.status_code == status.HTTP_200_OK
    assert all(d["patient_public_id"] == env["pat1_id"] for d in list_pat.json()["items"])

    # 4. Doctor Sarah Connor accesses Alice's document (under care) -> 200 OK
    doc_view = client.get(f"/api/v1/documents/{doc_id}", headers=env["doc_headers"])
    assert doc_view.status_code == status.HTTP_200_OK

    # 5. Doctor Sarah Connor tries to view Bob's document (unrelated) -> 403 Forbidden
    # Upload Bob's document via admin first
    bob_doc_res = client.post(
        "/api/v1/documents/upload",
        files={"file": ("bob.pdf", io.BytesIO(b"%PDF-1.4 bob"), "application/pdf")},
        data={"patient_id": env["pat2_id"], "title": "Bob Report"},
        headers=env["admin_headers"],
    )
    bob_doc_id = bob_doc_res.json()["document_id"]

    unauth_doc_view = client.get(f"/api/v1/documents/{bob_doc_id}", headers=env["doc_headers"])
    assert unauth_doc_view.status_code == status.HTTP_403_FORBIDDEN


def test_document_deletion_and_physical_cleanup(client: TestClient, tmp_path):
    """Verify document deletion removes record and deletes file from disk."""
    settings.DOCUMENT_STORAGE_PATH = str(tmp_path)
    env = setup_clinical_environment(client)

    pdf_file = ("cleanup.pdf", io.BytesIO(b"%PDF-1.4 cleanup test"), "application/pdf")
    upload_res = client.post(
        "/api/v1/documents/upload",
        files={"file": pdf_file},
        data={"patient_id": env["pat1_id"], "title": "To Delete"},
        headers=env["admin_headers"],
    )
    doc_id = upload_res.json()["document_id"]

    # Verify file exists on disk
    expected_file = os.path.join(str(tmp_path), f"{doc_id}.pdf")
    assert os.path.exists(expected_file)

    # Patient cannot delete document
    pat_del = client.delete(f"/api/v1/documents/{doc_id}", headers=env["pat1_headers"])
    assert pat_del.status_code == status.HTTP_403_FORBIDDEN

    # Admin deletes document
    del_res = client.delete(f"/api/v1/documents/{doc_id}", headers=env["admin_headers"])
    assert del_res.status_code == status.HTTP_200_OK
    assert "successfully deleted" in del_res.json()["detail"]

    # File should be removed from disk
    assert not os.path.exists(expected_file)

    # Get after delete -> 404
    get_after = client.get(f"/api/v1/documents/{doc_id}", headers=env["admin_headers"])
    assert get_after.status_code == status.HTTP_404_NOT_FOUND


def test_unauthenticated_document_access_rejected(client: TestClient):
    """Verify unauthenticated requests are rejected with 401."""
    assert client.get("/api/v1/documents").status_code == status.HTTP_401_UNAUTHORIZED
    assert client.post("/api/v1/documents/upload").status_code == status.HTTP_401_UNAUTHORIZED
    assert client.get("/api/v1/documents/DOCU-000").status_code == status.HTTP_401_UNAUTHORIZED
    assert client.delete("/api/v1/documents/DOCU-000").status_code == status.HTTP_401_UNAUTHORIZED

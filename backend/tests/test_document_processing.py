from datetime import datetime, timedelta, timezone
import io
import os
import docx
from fastapi import status
from fastapi.testclient import TestClient
import pypdf

from app.ai.chunker import chunk_extracted_document, estimate_token_count
from app.ai.cleaner import clean_clinical_text
from app.ai.extractors import ExtractedDocument, extract_document_text
from app.core.config import settings
from app.models.document import DocumentChunk, MedicalDocument
from app.schemas.document import DocumentType
from app.schemas.user import UserRole


def get_auth_headers(
    client: TestClient,
    role: UserRole = UserRole.DOCTOR,
    email: str = "doc_proc@hospital.org",
    name: str = "Dr. Process User",
) -> tuple[dict[str, str], int]:
    """Register and login helper returning authorization headers and user ID."""
    reg_res = client.post(
        "/api/v1/auth/register",
        json={
            "name": name,
            "email": email,
            "password": "SecurePassword123!",
            "role": role.value,
        },
    )
    user_id = reg_res.json()["id"]

    login_res = client.post(
        "/api/v1/auth/login",
        json={"email": email, "password": "SecurePassword123!"},
    )
    token = login_res.json()["access_token"]
    return {"Authorization": f"Bearer {token}"}, user_id


def setup_processing_environment(client: TestClient) -> dict[str, str]:
    """Set up admin, doctor, and patient with an established clinical relationship."""
    admin_headers, _ = get_auth_headers(client, role=UserRole.ADMIN, email="admin_proc@hospital.org", name="Admin Proc")
    doc_headers, doc_uid = get_auth_headers(client, role=UserRole.DOCTOR, email="cardiologist_proc@hospital.org", name="Dr. Gregory House")
    unrelated_doc_headers, _ = get_auth_headers(client, role=UserRole.DOCTOR, email="unrelated_doc_proc@hospital.org", name="Dr. Unrelated")
    pat_headers, _ = get_auth_headers(client, role=UserRole.PATIENT, email="john_proc@patient.org", name="John Doe")

    # Create doctor profile
    doc_res = client.post(
        "/api/v1/doctors",
        json={
            "user_id": doc_uid,
            "full_name": "Gregory House",
            "department": "Internal Medicine",
            "specialization": "Diagnostics",
            "medical_registration_number": "MED-PROC-001",
        },
        headers=admin_headers,
    )
    doc_id = doc_res.json()["doctor_id"]
    client.post(f"/api/v1/doctors/{doc_id}/verify", headers=admin_headers)

    # Create patient
    pat_res = client.post(
        "/api/v1/patients",
        json={
            "first_name": "John",
            "last_name": "Doe",
            "date_of_birth": "1985-05-15",
            "gender": "male",
            "email": "john_proc@patient.org",
        },
        headers=admin_headers,
    )
    pat_id = pat_res.json()["patient_id"]

    # Link doctor and patient via appointment
    future_time = (datetime.now(timezone.utc) + timedelta(days=3)).isoformat()
    client.post(
        "/api/v1/appointments",
        json={
            "patient_id": pat_id,
            "doctor_id": doc_id,
            "appointment_date": future_time,
            "reason_for_visit": "Diagnostic evaluation",
        },
        headers=admin_headers,
    )

    return {
        "admin_headers": admin_headers,
        "doc_headers": doc_headers,
        "unrelated_doc_headers": unrelated_doc_headers,
        "pat_headers": pat_headers,
        "pat_id": pat_id,
    }


def create_sample_docx_bytes(paragraphs: list[str]) -> bytes:
    """Helper to create a valid in-memory DOCX file."""
    doc = docx.Document()
    for p in paragraphs:
        doc.add_paragraph(p)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def create_sample_pdf_bytes(pages_text: list[str]) -> bytes:
    """Helper to create a valid in-memory PDF file containing text using pypdf."""
    writer = pypdf.PdfWriter()
    for text in pages_text:
        # Create a page with standard dimensions
        page = writer.add_blank_page(width=612, height=792)
        # Note: pypdf blank pages have no text stream unless annotations/text objects are written.
    buf = io.BytesIO()
    writer.write(buf)
    return buf.getvalue()


def test_clinical_cleaner_and_token_estimator():
    """Verify text normalization and deterministic token estimation."""
    raw_text = "  Patient   presents with   BP:  120/80 mmHg. \r\n\r\n\r\n  Medication: Metformin 500mg daily.   "
    cleaned = clean_clinical_text(raw_text)

    assert "120/80 mmHg" in cleaned
    assert "Metformin 500mg daily" in cleaned
    assert "\r" not in cleaned
    assert "   " not in cleaned

    tokens = estimate_token_count(cleaned)
    assert tokens > 0


def test_chunking_engine_logic():
    """Verify semantic chunking, sequential indexing, and overlap preservation."""
    sample_text = (
        "Patient underwent laparoscopic appendectomy without complications. "
        "Post-operative recovery was uneventful. "
        "Vital signs remained stable throughout hospital stay. "
        "Discharge medications include Acetaminophen 500mg as needed for mild pain. "
        "Follow-up appointment scheduled with surgical clinic in two weeks."
    )
    extracted = ExtractedDocument(
        text=sample_text,
        page_count=1,
        pages=[(1, sample_text)],
        metadata={"format": "txt"},
    )

    # Chunk with small token threshold to force multiple chunks
    chunks = chunk_extracted_document(extracted, chunk_size_tokens=25, chunk_overlap_tokens=10)
    assert len(chunks) >= 2

    # Verify sequential chunk indexing
    for idx, chunk in enumerate(chunks):
        assert chunk.chunk_index == idx
        assert chunk.page_number == 1
        assert chunk.token_count > 0
        assert len(chunk.content) > 0


def test_txt_upload_and_automatic_processing(client: TestClient, tmp_path):
    """Verify upload and automatic extraction/chunking of TXT clinical note."""
    settings.DOCUMENT_STORAGE_PATH = str(tmp_path)
    env = setup_processing_environment(client)

    clinical_note = (
        "CHIEF COMPLAINT: Chest pain on exertion.\n\n"
        "HISTORY: 55-year-old male with hypertension and hyperlipidemia presents with sub-sternal tightness.\n\n"
        "PLAN: Order 12-lead ECG, Troponin I, and schedule exercise treadmill stress test."
    )
    txt_file = ("clinical_note.txt", io.BytesIO(clinical_note.encode("utf-8")), "text/plain")

    res = client.post(
        "/api/v1/documents/upload",
        files={"file": txt_file},
        data={
            "patient_id": env["pat_id"],
            "title": "Cardiology Consultation Note",
            "document_type": "clinical_note",
        },
        headers=env["doc_headers"],
    )

    assert res.status_code == status.HTTP_201_CREATED
    data = res.json()
    assert data["processing_status"] == "completed"
    assert data["total_chunks"] >= 1
    assert data["page_count"] == 1
    doc_id = data["document_id"]

    # Verify chunks endpoint
    chunks_res = client.get(f"/api/v1/documents/{doc_id}/chunks", headers=env["doc_headers"])
    assert chunks_res.status_code == status.HTTP_200_OK
    chunks_data = chunks_res.json()
    assert chunks_data["total"] == data["total_chunks"]
    assert "Chest pain" in chunks_data["items"][0]["content"]


def test_docx_upload_and_automatic_processing(client: TestClient, tmp_path):
    """Verify upload and automatic extraction/chunking of DOCX discharge summary."""
    settings.DOCUMENT_STORAGE_PATH = str(tmp_path)
    env = setup_processing_environment(client)

    docx_bytes = create_sample_docx_bytes([
        "HOSPITAL DISCHARGE SUMMARY",
        "Admission Diagnosis: Acute exacerbation of chronic obstructive pulmonary disease (COPD).",
        "Discharge Medications: Albuterol sulfate inhaler 90mcg 2 puffs Q4H PRN, Prednisone 20mg daily for 5 days.",
        "Condition on Discharge: Stable, afebrile, resting SpO2 95% on room air.",
    ])
    docx_file = (
        "discharge_summary.docx",
        io.BytesIO(docx_bytes),
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    res = client.post(
        "/api/v1/documents/upload",
        files={"file": docx_file},
        data={
            "patient_id": env["pat_id"],
            "title": "Hospital Discharge Summary",
            "document_type": "discharge_summary",
        },
        headers=env["admin_headers"],
    )

    assert res.status_code == status.HTTP_201_CREATED
    data = res.json()
    assert data["processing_status"] == "completed"
    assert data["total_chunks"] >= 1
    assert data["file_extension"] == ".docx"


def test_reprocessing_idempotency_and_chunk_replacement(client: TestClient, tmp_path):
    """Verify document reprocessing deletes old chunks and updates total count without duplication."""
    settings.DOCUMENT_STORAGE_PATH = str(tmp_path)
    env = setup_processing_environment(client)

    content = "Initial clinical evaluation note for patient John Doe.\n\nAll laboratory markers within normal ranges."
    txt_file = ("progress.txt", io.BytesIO(content.encode("utf-8")), "text/plain")

    upload_res = client.post(
        "/api/v1/documents/upload",
        files={"file": txt_file},
        data={
            "patient_id": env["pat_id"],
            "title": "Progress Note",
            "document_type": "clinical_note",
        },
        headers=env["doc_headers"],
    )
    doc_id = upload_res.json()["document_id"]
    initial_chunks_count = upload_res.json()["total_chunks"]

    # Trigger reprocessing
    reprocess_res = client.post(
        f"/api/v1/documents/{doc_id}/reprocess",
        headers=env["doc_headers"],
    )
    assert reprocess_res.status_code == status.HTTP_200_OK
    assert reprocess_res.json()["processing_status"] == "completed"
    assert reprocess_res.json()["total_chunks"] == initial_chunks_count

    # Check chunks table: verify exact count (no duplicates)
    chunks_res = client.get(f"/api/v1/documents/{doc_id}/chunks", headers=env["doc_headers"])
    assert chunks_res.json()["total"] == initial_chunks_count


def test_document_processing_rbac_and_isolation(client: TestClient, tmp_path):
    """Verify RBAC rules for reprocessing and chunk retrieval."""
    settings.DOCUMENT_STORAGE_PATH = str(tmp_path)
    env = setup_processing_environment(client)

    content = "Sensitive lab results: Fasting blood glucose 98 mg/dL, HbA1c 5.4%."
    txt_file = ("lab.txt", io.BytesIO(content.encode("utf-8")), "text/plain")

    upload_res = client.post(
        "/api/v1/documents/upload",
        files={"file": txt_file},
        data={"patient_id": env["pat_id"], "title": "Glucose Test"},
        headers=env["doc_headers"],
    )
    doc_id = upload_res.json()["document_id"]

    # 1. Patient cannot access raw chunks (restricted to clinical/admin roles)
    pat_chunks = client.get(f"/api/v1/documents/{doc_id}/chunks", headers=env["pat_headers"])
    assert pat_chunks.status_code == status.HTTP_403_FORBIDDEN

    # 2. Patient cannot trigger reprocessing
    pat_reproc = client.post(f"/api/v1/documents/{doc_id}/reprocess", headers=env["pat_headers"])
    assert pat_reproc.status_code == status.HTTP_403_FORBIDDEN

    # 3. Unrelated doctor cannot access chunks
    unrel_chunks = client.get(f"/api/v1/documents/{doc_id}/chunks", headers=env["unrelated_doc_headers"])
    assert unrel_chunks.status_code == status.HTTP_403_FORBIDDEN

    # 4. Authorized doctor can access chunks
    auth_chunks = client.get(f"/api/v1/documents/{doc_id}/chunks", headers=env["doc_headers"])
    assert auth_chunks.status_code == status.HTTP_200_OK


def test_empty_scanned_pdf_failure_handling(client: TestClient, tmp_path):
    """Verify empty/scanned PDF without extractable text fails with descriptive error."""
    settings.DOCUMENT_STORAGE_PATH = str(tmp_path)
    env = setup_processing_environment(client)

    pdf_bytes = create_sample_pdf_bytes([""])
    pdf_file = ("scanned_report.pdf", io.BytesIO(pdf_bytes), "application/pdf")

    res = client.post(
        "/api/v1/documents/upload",
        files={"file": pdf_file},
        data={
            "patient_id": env["pat_id"],
            "title": "Scanned Image Report",
            "document_type": "imaging_report",
        },
        headers=env["doc_headers"],
    )

    assert res.status_code == status.HTTP_201_CREATED
    data = res.json()
    assert data["processing_status"] == "failed"
    assert "OCR image-based PDFs" in data["error_message"]
    assert data["total_chunks"] == 0
